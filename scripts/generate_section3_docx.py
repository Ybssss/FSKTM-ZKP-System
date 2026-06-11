from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


WORKSPACE_ROOT = Path(r"D:\STUDY\fyp project\fsktm-zkp-system")
PROJECT_ROOT = Path(
    r"D:\STUDY\s6\BIS20303KESELAMATAN WEB(SEM220252026) - khairulm[S1] - WebSec\lab\G1_S1_system"
)
APP_ROOT = PROJECT_ROOT / "fsktm-course-system"
OUTPUT_DIR = WORKSPACE_ROOT / "outputs" / "section3_report"
ASSET_DIR = OUTPUT_DIR / "screenshots"
DOCX_PATH = OUTPUT_DIR / "Section_3_Security_Testing_and_Demonstration.docx"


SECTIONS = [
    (
        "3.0 Security Testing and Demonstration",
        "This section presents the security testing carried out on the UTHM Secure Course Portal. "
        "The purpose of the testing was to verify that the implemented security controls worked "
        "correctly during registration, authentication, session handling, access control, data "
        "management, and audit logging. Testing was performed through source code review, database "
        "schema inspection, and direct observation of the implemented control flow in the supplied project files.",
    ),
    (
        "3.1 Registration Field Testing",
        "The registration page was tested by reviewing the required input fields and the related validation "
        "logic in the registration module. The system requires First Name, Last Name, User ID, Email, "
        "ID Number or Matric Number, Password, and Confirm Password before a registration can proceed.\n\n"
        "The validation logic confirms that incomplete submissions are rejected and that only valid input "
        "is accepted before account creation.",
    ),
    (
        "3.2 Data Sanitization Testing",
        "Registration and profile fields were tested against unsafe input patterns such as extra spaces, "
        "mixed case values, symbols, and HTML-like payloads. The system sanitizes User ID, email, phone, "
        "and matric values before processing them.\n\n"
        "This reduces the risk of unsafe values being stored or used directly in later operations.",
    ),
    (
        "3.3 Data Validation Testing",
        "Server-side validation was tested by reviewing the conditions used for invalid User IDs, email "
        "addresses, phone numbers, names, and date values. Invalid formats are rejected before database access.\n\n"
        "This confirms that the application does not rely only on browser-side validation.",
    ),
    (
        "3.4 Matric Year Blocking Testing",
        "The matric number rule was tested using values that represent different enrolment years. A value such "
        "as AI200159 represents the year 2020 and is rejected, while a newer value such as AI230159 is accepted "
        "when the format is correct.\n\n"
        "The system therefore enforces the project requirement that users enrolled in 2020 or earlier must be blocked.",
    ),
    (
        "3.5 Strong Password Testing",
        "Weak passwords such as lowercase-only, number-only, or passwords without special characters were "
        "checked against the password validation logic. The system requires at least eight characters, one "
        "uppercase letter, one lowercase letter, one number, and one special character.\n\n"
        "A password that does not satisfy all of these rules is rejected before registration can complete.",
    ),
    (
        "3.6 Suitable Error Message Testing",
        "Registration failures return specific validation messages that help the user fix the submission, such as "
        "invalid email format, invalid matric format, or blocked enrolment year.\n\n"
        "Login failures use a generic message stating that the User ID or Password is invalid. This prevents "
        "the login page from revealing whether a specific User ID exists.",
    ),
    (
        "3.7 Password Hash and Salt Testing",
        "The user table and registration flow were reviewed to confirm that passwords are never stored in plaintext. "
        "The registration process uses password_hash() with bcrypt-compatible output, and authentication uses "
        "password_verify() during login.\n\n"
        "Because bcrypt embeds a salt into the resulting hash, the stored password value is suitable for secure storage.",
    ),
    (
        "3.8 Google reCAPTCHA Testing",
        "The login module includes Google reCAPTCHA verification before credential checking is allowed to proceed. "
        "If the CAPTCHA response is missing or invalid, the login request is rejected and an appropriate message is shown.\n\n"
        "This provides additional protection against automated login attempts.",
    ),
    (
        "3.9 SQL Injection Testing",
        "The login, registration, user management, and course management flows were reviewed for SQL injection protection. "
        "These modules use prepared statements with parameter binding instead of concatenating raw user input into SQL strings.\n\n"
        "This means common payloads such as ' OR '1'='1 are treated as data rather than executable SQL.",
    ),
    (
        "3.10 Cross-Site Scripting Testing",
        "Display points for user and course data were reviewed for output encoding. The system uses htmlspecialchars() "
        "through helper functions such as e() and ucp_e() before rendering values back into HTML.\n\n"
        "This prevents stored or reflected script payloads from executing in the browser when unsafe input is displayed.",
    ),
    (
        "3.11 Account Lockout Testing",
        "The login logic records failed attempts for existing accounts by updating failed_login_attempts and locked_until "
        "in the users table. Once the maximum failed-attempt threshold is reached, the account is temporarily locked.\n\n"
        "This reduces the risk of repeated password guessing against a known account.",
    ),
    (
        "3.12 Browser / Client Lockout Testing",
        "The system also tracks repeated failed login attempts from the same browser or client using a hashed client key "
        "stored in the login_attempt_locks table. When the threshold is reached, further login attempts from that client "
        "are blocked until the lock expires.\n\n"
        "This protection still applies even when the attacker uses non-existing usernames.",
    ),
    (
        "3.13 Session Management Testing",
        "After a successful login, the application stores the authenticated session and updates the last activity timestamp. "
        "Protected pages check whether a valid session exists before granting access.\n\n"
        "The session ID is regenerated after successful authentication, which helps reduce session fixation risk.",
    ),
    (
        "3.14 Session Timeout Testing",
        "The system defines a session inactivity timeout through SESSION_TIMEOUT_SECONDS. When the inactivity period is exceeded, "
        "the active session is destroyed and the user is redirected to the login page with a session-expired warning.\n\n"
        "This prevents old sessions from remaining active indefinitely.",
    ),
    (
        "3.15 Logout and Session Destruction Testing",
        "The logout module clears the session data, invalidates the session cookie, destroys the session, and redirects the user "
        "to the login page. Cache-control headers are also sent to reduce the chance of protected pages being restored from browser history.\n\n"
        "This confirms that logout closes the authenticated session cleanly.",
    ),
    (
        "3.16 Broken Access Control Testing",
        "Role-based access checks were reviewed across admin, student, and lecturer routes. Admin-only pages and APIs check both "
        "authentication state and user role before processing requests.\n\n"
        "This prevents a normal student account from browsing directly to administrator functions or calling privileged APIs.",
    ),
    (
        "3.17 HTTPS / SSL Testing",
        "The .htaccess configuration includes HTTP-to-HTTPS redirection and multiple browser security headers such as "
        "X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy, and Content-Security-Policy.\n\n"
        "These settings improve transport security and reduce several browser-based attack surfaces.",
    ),
    (
        "3.18 Audit Log Testing",
        "The project records security-relevant events in the audit_logs table through the logAuditEvent() helper. Logged events include "
        "login success, login failure, CAPTCHA failure, lockout, session timeout, and administrator actions.\n\n"
        "This provides traceability for security monitoring and later investigation.",
    ),
    (
        "3.19 Admin User Management Testing",
        "The administrator user-management module was reviewed for validation, authorization, password reset handling, self-protection "
        "controls, and audit logging. The module validates user data server-side, prevents an administrator from deleting or deactivating "
        "their own account, and hashes temporary passwords before storage.\n\n"
        "This confirms that the user-management feature follows the same security rules as public registration and login.",
    ),
    (
        "3.20 Course Management Testing",
        "The administrator course-management module was reviewed for duplicate-code checking, prepared statements, authorization checks, "
        "and audit logging. Adding and updating courses is done through prepared statements, while deletion is restricted when active "
        "registrations already exist.\n\n"
        "This confirms that the course module protects both data integrity and backend query safety.",
    ),
]


FIGURES = [
    {
        "filename": "figure_3_1_registration_validation.png",
        "title": "Figure 3.1 Registration validation and matric-year blocking",
        "caption": (
            "Evidence from register.php showing User ID validation, email validation, "
            "matric format validation, and the rule that blocks users enrolled in 2020 or earlier."
        ),
        "segments": [
            {
                "path": APP_ROOT / "pages" / "auth" / "register.php",
                "start": "// User ID validation",
                "end": "// Optional field validation",
            }
        ],
    },
    {
        "filename": "figure_3_2_password_security.png",
        "title": "Figure 3.2 Password policy and hashed storage",
        "caption": (
            "Evidence from functions.php, register.php, and schema.sql showing the password complexity rules, "
            "bcrypt-based password hashing, and the database password column used for hashed storage."
        ),
        "segments": [
            {
                "path": APP_ROOT / "includes" / "functions.php",
                "start": "if (!function_exists('getPasswordValidationErrors')) {",
                "end": "if (!function_exists('logAuditEvent')) {",
            },
            {
                "path": APP_ROOT / "pages" / "auth" / "register.php",
                "start": "$hashed_password = password_hash($password, PASSWORD_BCRYPT);",
                "lines_after": 16,
            },
            {
                "path": APP_ROOT / "database" / "schema.sql",
                "start": "CREATE TABLE users (",
                "end": ") ENGINE=InnoDB;",
            },
        ],
    },
    {
        "filename": "figure_3_3_login_protection.png",
        "title": "Figure 3.3 Login protection, CAPTCHA, and lockout flow",
        "caption": (
            "Evidence from login.php showing browser/client lockout checks, session-based rate limiting, "
            "Google reCAPTCHA verification, password_verify(), and session_regenerate_id(true)."
        ),
        "segments": [
            {
                "path": APP_ROOT / "pages" / "auth" / "login.php",
                "start": "$clientKeyHash = getLoginClientKeyHash();",
                "end": "$stmt = $conn->prepare(\"SELECT user_id, username, email, password, first_name, last_name, user_type, profile_picture, is_active, failed_login_attempts, locked_until FROM users WHERE username = ? LIMIT 1\");",
            },
            {
                "path": APP_ROOT / "pages" / "auth" / "login.php",
                "start": "} elseif ($user && (int)$user['is_active'] === 1 && password_verify($password, $user['password'])) {",
                "lines_after": 24,
            },
        ],
    },
    {
        "filename": "figure_3_4_session_access_control.png",
        "title": "Figure 3.4 Session timeout and role-based access control",
        "caption": (
            "Evidence from functions.php showing session timeout handling, role helper functions, "
            "and the reusable administrator access requirement used by privileged pages."
        ),
        "segments": [
            {
                "path": APP_ROOT / "includes" / "functions.php",
                "start": "function isLoggedIn() {",
                "end": "function sanitizeInput($data) {",
            },
            {
                "path": APP_ROOT / "includes" / "functions.php",
                "start": "if (!function_exists('requireAdminAccess')) {",
                "end": "if (!function_exists('sanitizePersonName')) {",
            },
        ],
    },
    {
        "filename": "figure_3_5_https_headers.png",
        "title": "Figure 3.5 HTTPS redirection and security headers",
        "caption": (
            "Evidence from .htaccess showing HTTP-to-HTTPS redirection and the configured browser security headers, "
            "including X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy, and Content-Security-Policy."
        ),
        "segments": [
            {
                "path": APP_ROOT / ".htaccess",
                "start": "RewriteEngine On",
                "end": "RewriteRule ^(.*)$ index.php [QSA,L]",
            }
        ],
    },
    {
        "filename": "figure_3_6_audit_and_lock_tables.png",
        "title": "Figure 3.6 Audit log and client lock database evidence",
        "caption": (
            "Evidence from fsktm_courses.sql showing the audit_logs table and the login_attempt_locks table "
            "used for event recording and browser or client lockout tracking."
        ),
        "segments": [
            {
                "path": PROJECT_ROOT / "fsktm_courses.sql",
                "start": "CREATE TABLE `audit_logs` (",
                "end": ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;",
            },
            {
                "path": PROJECT_ROOT / "fsktm_courses.sql",
                "start": "CREATE TABLE `login_attempt_locks` (",
                "end": ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;",
            },
        ],
    },
    {
        "filename": "figure_3_7_admin_user_controls.png",
        "title": "Figure 3.7 Administrator user-management protection",
        "caption": (
            "Evidence from manage-users.php and user-api.php showing self-delete and self-deactivate protection, "
            "hashed temporary passwords, and audit logging for administrator-driven user actions."
        ),
        "segments": [
            {
                "path": APP_ROOT / "pages" / "admin" / "manage-users.php",
                "start": "function ucp_toggle_user_status($conn, $userId) {",
                "end": "function ucp_delete_user($conn, $userId) {",
            },
            {
                "path": APP_ROOT / "pages" / "admin" / "manage-users.php",
                "start": "function ucp_delete_user($conn, $userId) {",
                "end": "$db = new Database();",
            },
            {
                "path": APP_ROOT / "api" / "user-api.php",
                "start": "function handleResetUserPassword(mysqli $conn): array {",
                "end": "if (!function_exists('generateStaffId')) {",
            },
        ],
    },
    {
        "filename": "figure_3_8_course_prepared_statements.png",
        "title": "Figure 3.8 Course-management prepared statements",
        "caption": (
            "Evidence from course-api.php showing duplicate-course checks, prepared statements, "
            "parameter binding, and audit logging for add and update course operations."
        ),
        "segments": [
            {
                "path": APP_ROOT / "api" / "course-api.php",
                "start": "function handleAddCourse($conn) {",
                "end": "function handleUpdateCourse($conn) {",
            },
            {
                "path": APP_ROOT / "api" / "course-api.php",
                "start": "function handleUpdateCourse($conn) {",
                "end": "?>",
            },
        ],
    },
]


def extract_segment(path: Path, start: str, end: str | None = None, lines_after: int | None = None) -> tuple[str, int]:
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    start_index = None
    for index, line in enumerate(lines):
        if start in line:
            start_index = index
            break

    if start_index is None:
        raise ValueError(f"Start marker not found in {path}: {start}")

    if end is not None:
        end_index = None
        for index in range(start_index, len(lines)):
            if end in lines[index]:
                end_index = index
                break
        if end_index is None:
            raise ValueError(f"End marker not found in {path}: {end}")
        selected = lines[start_index : end_index + 1]
    elif lines_after is not None:
        selected = lines[start_index : start_index + lines_after + 1]
    else:
        selected = lines[start_index : start_index + 24]

    numbered = [f"{start_index + idx + 1:>4} | {line}" for idx, line in enumerate(selected)]
    return "\n".join(numbered), start_index + 1


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path(r"C:\Windows\Fonts\consola.ttf"),
        Path(r"C:\Windows\Fonts\cour.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def wrap_line(line: str, width: int = 118) -> list[str]:
    if len(line) <= width:
        return [line]

    wrapped: list[str] = []
    remaining = line
    continuation = " " * 7
    while len(remaining) > width:
        wrapped.append(remaining[:width])
        remaining = continuation + remaining[width:]
    wrapped.append(remaining)
    return wrapped


def render_figure(figure: dict) -> Path:
    title = figure["title"]
    output_path = ASSET_DIR / figure["filename"]
    body_lines = []

    for segment in figure["segments"]:
        text, _ = extract_segment(
            path=segment["path"],
            start=segment["start"],
            end=segment.get("end"),
            lines_after=segment.get("lines_after"),
        )
        body_lines.append(f"[File] {segment['path']}")
        body_lines.extend(text.splitlines())
        body_lines.append("")

    code_lines: list[str] = []
    for line in body_lines:
        code_lines.extend(wrap_line(line))

    title_font = load_font(30)
    code_font = load_font(22)
    small_font = load_font(18)

    line_height = code_font.getbbox("Ag")[3] + 10
    code_height = max(10, len(code_lines)) * line_height
    width = 1650
    margin = 40
    header_height = 110
    footer_height = 44
    height = header_height + code_height + footer_height + margin

    image = Image.new("RGB", (width, height), (246, 248, 250))
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((20, 20, width - 20, height - 20), radius=28, fill=(255, 255, 255), outline=(220, 224, 230))
    draw.rectangle((20, 20, width - 20, header_height + 10), fill=(30, 41, 59))
    draw.text((50, 45), title, fill=(255, 255, 255), font=title_font)
    draw.text((50, 82), "Source-code and database evidence captured from the supplied project files", fill=(203, 213, 225), font=small_font)

    top = header_height + 34
    for line in code_lines:
        if line.startswith("[File] "):
            draw.text((50, top), line, fill=(13, 110, 253), font=small_font)
            top += line_height
            continue
        draw.text((50, top), line, fill=(31, 41, 55), font=code_font)
        top += line_height

    image.save(output_path)
    return output_path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx(figure_paths: list[tuple[str, Path, str]]) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Section 3: Security Testing and Demonstration")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("UTHM Secure Course Portal")
    subtitle_run.italic = True
    subtitle_run.font.name = "Times New Roman"

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Note: The screenshots embedded in this section are evidence images generated from the provided source files "
        "and database schema because the project was supplied as source code rather than as a running deployed environment."
    )
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(89, 89, 89)

    figure_map = {
        "3.6 Suitable Error Message Testing": ["Figure 3.1 Registration validation and matric-year blocking"],
        "3.12 Browser / Client Lockout Testing": [
            "Figure 3.2 Password policy and hashed storage",
            "Figure 3.3 Login protection, CAPTCHA, and lockout flow",
        ],
        "3.16 Broken Access Control Testing": ["Figure 3.4 Session timeout and role-based access control"],
        "3.17 HTTPS / SSL Testing": ["Figure 3.5 HTTPS redirection and security headers"],
        "3.18 Audit Log Testing": ["Figure 3.6 Audit log and client lock database evidence"],
        "3.19 Admin User Management Testing": ["Figure 3.7 Administrator user-management protection"],
        "3.20 Course Management Testing": ["Figure 3.8 Course-management prepared statements"],
    }

    figure_lookup = {title: (path, caption) for title, path, caption in figure_paths}

    for heading_text, body_text in SECTIONS:
        heading = doc.add_paragraph()
        heading.style = doc.styles["Heading 2"]
        heading_run = heading.add_run(heading_text)
        heading_run.font.name = "Times New Roman"

        for paragraph_text in body_text.split("\n\n"):
            para = doc.add_paragraph(paragraph_text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if heading_text in figure_map:
            for figure_title in figure_map[heading_text]:
                image_path, caption = figure_lookup[figure_title]
                doc.add_picture(str(image_path), width=Inches(6.5))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(caption)
                cap_run.italic = True
                cap_run.font.name = "Times New Roman"

    appendix = doc.add_paragraph()
    appendix.style = doc.styles["Heading 2"]
    appendix_run = appendix.add_run("Screenshot List")
    appendix_run.font.name = "Times New Roman"

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Figure"
    header_cells[1].text = "Purpose"
    set_cell_shading(header_cells[0], "D9EAF7")
    set_cell_shading(header_cells[1], "D9EAF7")

    for figure_title, _, caption in figure_paths:
        row_cells = table.add_row().cells
        row_cells[0].text = figure_title.split(" ", 2)[0] + " " + figure_title.split(" ", 2)[1]
        row_cells[1].text = caption

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    figure_paths: list[tuple[str, Path, str]] = []
    for figure in FIGURES:
        path = render_figure(figure)
        figure_paths.append((figure["title"], path, figure["caption"]))

    build_docx(figure_paths)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
